import os
import asyncio
import logging
from typing import List, Optional, Dict, Any, Tuple
from PIL import Image
from dataclasses import dataclass
from datetime import datetime
import json
import tempfile
import requests
import shutil
import threading
import time
from pdf2image import convert_from_path
import magic

import google.generativeai as genai
from dotenv import load_dotenv
from utils.pydantic_objectid import PyObjectId
from models.pdf import PDF, ProcessingStatus
from models.table import Table

@dataclass
class ExtractedTable:
    table_id: int
    title: str
    markdown_content: str
    column_headers: List[str]
    row_count: int
    column_count: int
    merge_start_page: Optional[int] = None  # Track where merging started

@dataclass
class PageResponse:
    page_number: int
    raw_response: str
    tables: List[ExtractedTable]

class BackgroundTableExtractor:
    """
    BULLETPROOF TWO-PHASE PIPELINE - NO MORE FUCKUPS
    """
    
    def __init__(self):
        self.logger = self._setup_logger()
        self.api_keys = self._load_api_keys()
        
        if not self.api_keys:
            raise ValueError("No Gemini API keys found")
        
        self.clients = [genai.Client(api_key=key) for key in self.api_keys]
        self.client_index = 0
        self.semaphore = asyncio.Semaphore(min(30, len(self.api_keys)))
        self.logger.info(f"BackgroundTableExtractor initialized - BULLETPROOF PIPELINE with {len(self.api_keys)} keys")

    def _setup_logger(self) -> logging.Logger:
        """Setup logger"""
        logger = logging.getLogger("BackgroundTableExtractor")
        for handler in logger.handlers[:]:
            logger.removeHandler(handler)
        
        handler = logging.StreamHandler()
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        logger.addHandler(handler)
        logger.setLevel(logging.INFO)
        return logger

    def _load_api_keys(self) -> List[str]:
        """Load API keys"""
        load_dotenv()
        keys_string = os.getenv("GEMINI_KEYS", "")
        if not keys_string.strip():
            return []
        return [k.strip() for k in keys_string.split(",") if k.strip()]

    def _get_next_client(self):
        """Thread-safe client rotation"""
        current_thread_id = threading.get_ident()
        client_idx = (current_thread_id + int(time.time() * 1000)) % len(self.clients)
        return self.clients[client_idx]

    async def extract_tables_for_pdf(self, pdf_id: str) -> Dict[str, Any]:
        """MAIN METHOD: Bulletproof two-phase pipeline"""
        start_time = datetime.now()
        
        try:
            pdf_record = await PDF.get(pdf_id)
            if not pdf_record:
                return {"success": False, "error": "PDF not found"}
            
            pdf_record.processing_status = ProcessingStatus.BACKGROUND_PROCESSING
            await pdf_record.save()
            
            self.logger.info(f"🚀 Starting BULLETPROOF TWO-PHASE extraction for: {pdf_record.filename}")
            
            # Generate page images
            images_folder = await self._generate_page_images(pdf_record)
            if not images_folder:
                raise Exception("Failed to generate page images")
            
            # PHASE 1: PARALLEL EXTRACTION
            self.logger.info(f"⚡ PHASE 1: PARALLEL extraction...")
            phase1_start = datetime.now()
            page_responses = await self._phase1_parallel_extraction(pdf_record, images_folder)
            phase1_time = (datetime.now() - phase1_start).total_seconds()
            self.logger.info(f"✅ PHASE 1 completed in {phase1_time:.2f}s - processed {len(page_responses)} pages")
            
            # PHASE 2: BULLETPROOF SEQUENTIAL MERGING
            self.logger.info(f"🔗 PHASE 2: BULLETPROOF sequential merging...")
            phase2_start = datetime.now()
            total_tables = await self._phase2_bulletproof_sequential_merging(pdf_record, page_responses)
            phase2_time = (datetime.now() - phase2_start).total_seconds()
            self.logger.info(f"✅ PHASE 2 completed in {phase2_time:.2f}s - inserted {total_tables} tables")
            
            # Update completion
            pdf_record.processing_status = ProcessingStatus.COMPLETED
            pdf_record.total_tables_found = total_tables
            pdf_record.tables_processed = total_tables
            pdf_record.fully_completed_at = datetime.utcnow()
            await pdf_record.save()
            
            # Cleanup
            await self._cleanup_temp_images(images_folder)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            self.logger.info(f"🎯 BULLETPROOF extraction completed!")
            self.logger.info(f"📊 Results: {total_tables} tables in {processing_time:.2f}s")
            
            return {
                "success": True,
                "tables_extracted": total_tables,
                "processing_time": processing_time,
                "phase1_time": phase1_time,
                "phase2_time": phase2_time,
                "status": "Bulletproof extraction completed successfully"
            }
            
        except Exception as e:
            self.logger.error(f"❌ Bulletproof extraction failed: {e}")
            
            if 'pdf_record' in locals():
                pdf_record.processing_status = ProcessingStatus.FAILED
                pdf_record.background_error = str(e)
                await pdf_record.save()
            
            return {"success": False, "error": str(e)}

    async def _generate_page_images(self, pdf_record: PDF) -> Optional[str]:
        """Generate page images with AUTOMATIC WORD-TO-PDF CONVERSION"""
        try:
            from pdf2image import convert_from_path
            
            temp_folder = tempfile.mkdtemp(prefix=f"bulletproof_extraction_{pdf_record.id}_")
            
            self.logger.info(f"🖼️ Generating images for {pdf_record.page_count} pages...")
            
            # Download file
            response = requests.get(pdf_record.cloudinary_url, timeout=120)
            response.raise_for_status()
            
            self.logger.info(f"✅ Downloaded file ({len(response.content)/1024/1024:.1f}MB)")
            
            # Save file temporarily
            temp_file_path = os.path.join(temp_folder, f"temp_{pdf_record.id}")
            with open(temp_file_path, 'wb') as f:
                f.write(response.content)
            
            # Detect file type
            file_type = self._detect_file_type(temp_file_path)
            self.logger.info(f"📄 Detected file type: {file_type}")
            
            # Handle different file types
            if file_type == 'application/pdf':
                # Process PDF directly
                return await self._process_pdf_file(temp_file_path, temp_folder)
                
            elif file_type in ['image/png', 'image/jpeg', 'image/jpg']:
                # Process image
                return await self._process_image_file(temp_file_path, temp_folder)
                
            elif file_type in [
                'application/msword', 
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/vnd.ms-word'
            ]:
                # 🚀 AUTOMATIC WORD-TO-PDF CONVERSION
                self.logger.info(f"📝 Converting Word document to PDF...")
                pdf_path = await self._convert_word_to_pdf(temp_file_path, temp_folder)
                return await self._process_pdf_file(pdf_path, temp_folder)
                
            elif file_type in [
                'application/vnd.ms-excel',
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            ]:
                # Excel documents - clear error  
                self.logger.error(f"❌ Excel documents are not supported")
                raise Exception("Excel documents are not supported. Please convert your document to PDF format first and try again.")
                
            else:
                # Unknown file type
                self.logger.error(f"❌ Unsupported file type: {file_type}")
                raise Exception(f"Unsupported file type: {file_type}. Please upload a PDF, Word document, or image file.")
                
        except Exception as e:
            self.logger.error(f"❌ Error generating images: {e}")
            
            # Cleanup on error
            if 'temp_folder' in locals() and os.path.exists(temp_folder):
                try:
                    shutil.rmtree(temp_folder)
                except:
                    pass
            
            return None

    async def _convert_word_to_pdf(self, word_file_path: str, temp_folder: str) -> str:
        """🚀 AUTOMATIC WORD-TO-PDF CONVERSION"""
        try:
            # Method 1: Try python-docx2pdf (fastest)
            try:
                from docx2pdf import convert
                
                pdf_path = os.path.join(temp_folder, f"converted_{int(time.time())}.pdf")
                
                self.logger.info(f"📝 Converting Word to PDF using docx2pdf...")
                await asyncio.get_event_loop().run_in_executor(
                    None, 
                    lambda: convert(word_file_path, pdf_path)
                )
                
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    self.logger.info(f"✅ Word-to-PDF conversion successful (docx2pdf)")
                    return pdf_path
                    
            except ImportError:
                self.logger.info(f"📝 docx2pdf not available, trying alternative method...")
            except Exception as e:
                self.logger.warning(f"⚠️ docx2pdf conversion failed: {e}")
            
            # Method 2: Try python-docx + reportlab (more reliable)
            try:
                from docx import Document
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib import colors
                
                self.logger.info(f"📝 Converting Word to PDF using docx + reportlab...")
                
                # Read Word document
                doc = Document(word_file_path)
                
                # Create PDF
                pdf_path = os.path.join(temp_folder, f"converted_{int(time.time())}.pdf")
                pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                
                styles = getSampleStyleSheet()
                story = []
                
                # Extract content from Word document
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        story.append(Paragraph(paragraph.text, styles['Normal']))
                        story.append(Spacer(1, 12))
                
                # Extract tables from Word document
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    
                    if table_data:
                        pdf_table = Table(table_data)
                        pdf_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 14),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(pdf_table)
                        story.append(Spacer(1, 12))
                
                # Build PDF
                await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: pdf_doc.build(story)
                )
                
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    self.logger.info(f"✅ Word-to-PDF conversion successful (docx + reportlab)")
                    return pdf_path
                    
            except ImportError:
                self.logger.warning(f"⚠️ python-docx or reportlab not available")
            except Exception as e:
                self.logger.warning(f"⚠️ docx + reportlab conversion failed: {e}")
            
            # Method 3: Try LibreOffice (if available)
            try:
                import subprocess
                
                self.logger.info(f"📝 Converting Word to PDF using LibreOffice...")
                
                pdf_path = os.path.join(temp_folder, f"converted_{int(time.time())}.pdf")
                
                # Try LibreOffice conversion
                result = await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', temp_folder, word_file_path
                    ], capture_output=True, timeout=60)
                )
                
                # Find the generated PDF
                for file in os.listdir(temp_folder):
                    if file.endswith('.pdf'):
                        generated_pdf = os.path.join(temp_folder, file)
                        if os.path.getsize(generated_pdf) > 0:
                            self.logger.info(f"✅ Word-to-PDF conversion successful (LibreOffice)")
                            return generated_pdf
                            
            except Exception as e:
                self.logger.warning(f"⚠️ LibreOffice conversion failed: {e}")
            
            # All methods failed
            raise Exception("Could not convert Word document to PDF. Please install one of: docx2pdf, python-docx+reportlab, or LibreOffice")
            
        except Exception as e:
            self.logger.error(f"❌ Word-to-PDF conversion failed: {e}")
            raise Exception(f"Failed to convert Word document: {e}")

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type reliably"""
        try:
            # Try python-magic first
            try:
                import magic
                if hasattr(magic, 'from_file'):
                    return magic.from_file(file_path, mime=True)
            except ImportError:
                pass
            
            # Fallback to header detection
            with open(file_path, 'rb') as f:
                header = f.read(16)
                
                if header.startswith(b'%PDF'):
                    return 'application/pdf'
                elif header.startswith(b'\x89PNG'):
                    return 'image/png'
                elif header.startswith(b'\xff\xd8\xff'):
                    return 'image/jpeg'
                elif header.startswith(b'PK\x03\x04'):
                    # ZIP-based format
                    with open(file_path, 'rb') as zf:
                        content = zf.read(1024)
                        if b'word/' in content:
                            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        elif b'xl/' in content:
                            return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        else:
                            return 'application/zip'
                elif header.startswith(b'\xd0\xcf\x11\xe0'):
                    return 'application/msword'
                else:
                    return 'unknown'
                    
        except Exception as e:
            self.logger.warning(f"⚠️ File type detection error: {e}")
            return 'unknown'

    async def _process_pdf_file(self, file_path: str, temp_folder: str) -> str:
        """Process PDF file"""
        try:
            # Ensure PDF extension
            if not file_path.endswith('.pdf'):
                pdf_path = file_path + '.pdf'
                if file_path != pdf_path:
                    os.rename(file_path, pdf_path)
                file_path = pdf_path
            
            # Validate PDF
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                actual_page_count = len(reader.pages)
                self.logger.info(f"📄 PDF validation successful: {actual_page_count} pages")
            except Exception as pdf_error:
                self.logger.error(f"❌ Invalid PDF file: {pdf_error}")
                raise Exception(f"File appears to be corrupted or not a valid PDF: {pdf_error}")
            
            # Convert to images
            images = convert_from_path(file_path, dpi=200, fmt='PNG')
            
            for i, image in enumerate(images):
                page_num = i + 1
                image_path = os.path.join(temp_folder, f"page_{page_num:03d}.png")
                image.save(image_path, 'PNG')
            
            self.logger.info(f"✅ Generated {len(images)} page images")
            
            return temp_folder
            
        except Exception as e:
            raise Exception(f"Failed to process PDF file: {e}")

    async def _process_image_file(self, file_path: str, temp_folder: str) -> str:
        """Process single image file"""
        try:
            from PIL import Image
            
            image = Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            image_path = os.path.join(temp_folder, "page_001.png")
            image.save(image_path, 'PNG')
            
            self.logger.info(f"✅ Generated 1 page image from source image")
            
            return temp_folder
            
        except Exception as e:
            raise Exception(f"Failed to process image file: {e}")


    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type reliably"""
        try:
            # Try python-magic first
            try:
                import magic
                if hasattr(magic, 'from_file'):
                    return magic.from_file(file_path, mime=True)
            except ImportError:
                pass
            
            # Fallback to header detection
            with open(file_path, 'rb') as f:
                header = f.read(16)
                
                if header.startswith(b'%PDF'):
                    return 'application/pdf'
                elif header.startswith(b'\x89PNG'):
                    return 'image/png'
                elif header.startswith(b'\xff\xd8\xff'):
                    return 'image/jpeg'
                elif header.startswith(b'PK\x03\x04'):
                    # ZIP-based format (could be docx, xlsx, etc.)
                    with open(file_path, 'rb') as zf:
                        content = zf.read(1024)
                        if b'word/' in content:
                            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        elif b'xl/' in content:
                            return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        else:
                            return 'application/zip'
                elif header.startswith(b'\xd0\xcf\x11\xe0'):
                    # Old Office format
                    return 'application/msword'
                else:
                    return 'unknown'
                    
        except Exception as e:
            self.logger.warning(f"⚠️ File type detection error: {e}")
            return 'unknown'

    async def _process_pdf_file(self, file_path: str, temp_folder: str) -> str:
        """Process PDF file"""
        try:
            # Rename to PDF extension
            pdf_path = file_path + '.pdf'
            os.rename(file_path, pdf_path)
            
            # Validate PDF
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(pdf_path)
                actual_page_count = len(reader.pages)
                self.logger.info(f"📄 PDF validation successful: {actual_page_count} pages")
            except Exception as pdf_error:
                self.logger.error(f"❌ Invalid PDF file: {pdf_error}")
                raise Exception(f"File appears to be corrupted or not a valid PDF: {pdf_error}")
            
            # Convert to images
            images = convert_from_path(pdf_path, dpi=200, fmt='PNG')
            
            for i, image in enumerate(images):
                page_num = i + 1
                image_path = os.path.join(temp_folder, f"page_{page_num:03d}.png")
                image.save(image_path, 'PNG')
            
            self.logger.info(f"✅ Generated {len(images)} page images from PDF")
            
            # Cleanup
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            
            return temp_folder
            
        except Exception as e:
            raise Exception(f"Failed to process PDF file: {e}")

    async def _process_image_file(self, file_path: str, temp_folder: str) -> str:
        """Process single image file"""
        try:
            from PIL import Image
            
            image = Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            image_path = os.path.join(temp_folder, "page_001.png")
            image.save(image_path, 'PNG')
            
            self.logger.info(f"✅ Generated 1 page image from source image")
            
            # Cleanup
            if os.path.exists(file_path):
                os.remove(file_path)
            
            return temp_folder
            
        except Exception as e:
            raise Exception(f"Failed to process image file: {e}")



    async def _phase1_parallel_extraction(self, pdf_record: PDF, images_folder: str) -> List[PageResponse]:
        """PHASE 1: PARALLEL extraction"""
        
        tasks = []
        for page_num in range(1, pdf_record.page_count + 1):
            image_path = os.path.join(images_folder, f"page_{page_num:03d}.png")
            if os.path.exists(image_path):
                task = self._extract_from_single_page(image_path, page_num)
                tasks.append(task)
        
        self.logger.info(f"⚡ PHASE 1: Processing {len(tasks)} pages in PARALLEL...")
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        page_responses = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                self.logger.error(f"❌ PHASE 1: Page {i+1} failed: {result}")
            elif result:
                page_responses.append(result)
                self.logger.debug(f"✅ PHASE 1: Page {result.page_number} - {len(result.tables)} tables")
        
        return page_responses

    async def _extract_from_single_page(self, image_path: str, page_num: int) -> Optional[PageResponse]:
        """Extract all tables from a single page"""
        async with self.semaphore:
            try:
                image = Image.open(image_path)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                prompt = f"""
EXTRACT ALL TABLES - Page {page_num}

Extract ALL tables from this page. Return them in this exact format:

If tables exist:
{{Table 1: descriptive_title}}
| header1 | header2 | header3 |
|---------|---------|---------|
| data1   | data2   | data3   |
| data4   | data5   | data6   |

{{Table 2: another_title}}
| col1 | col2 |
|------|------|
| val1 | val2 |

If no tables: return exactly "EMPTY"

CRITICAL RULES:
- Extract REAL data, not placeholders
- Include ALL visible data in proper markdown format
- Use descriptive titles
- Each table must be complete with headers and data
- If no tables exist: return "EMPTY"

Extract now:
"""
                
                for attempt in range(2):
                    try:
                        client = self._get_next_client()
                        
                        response = await asyncio.wait_for(
                            asyncio.get_event_loop().run_in_executor(
                                None,
                                lambda: client.models.generate_content(
                                    model="gemini-2.5-flash-preview-04-17",
                                    contents=[image, prompt]
                                )
                            ),
                            timeout=60.0
                        )
                        
                        if response and hasattr(response, 'text') and response.text:
                            return self._parse_page_response(str(response.text), page_num)
                        
                    except Exception as e:
                        if attempt == 1:
                            self.logger.warning(f"⚠️ PHASE 1: Page {page_num} failed: {e}")
                            break
                        await asyncio.sleep(1)
                
                return PageResponse(page_num, "EMPTY", [])
                
            except Exception as e:
                self.logger.error(f"❌ PHASE 1: Page {page_num} error: {e}")
                return PageResponse(page_num, "EMPTY", [])

    def _parse_page_response(self, response_text: str, page_num: int) -> PageResponse:
        """Parse page response into individual tables"""
        try:
            if response_text.strip().upper() == "EMPTY":
                return PageResponse(page_num, response_text, [])
            
            tables = []
            current_table_content = ""
            current_title = ""
            table_id = 1
            
            lines = response_text.split('\n')
            
            for line in lines:
                if line.strip().startswith('{Table') and '}' in line:
                    # Save previous table if exists
                    if current_table_content.strip():
                        table = self._create_table_from_content(current_title, current_table_content.strip(), table_id, page_num)
                        if table:
                            tables.append(table)
                            table_id += 1
                    
                    # Start new table
                    current_title = line.strip()
                    current_table_content = ""
                else:
                    current_table_content += line + '\n'
            
            # Save last table
            if current_table_content.strip():
                table = self._create_table_from_content(current_title, current_table_content.strip(), table_id, page_num)
                if table:
                    tables.append(table)
            
            return PageResponse(page_num, response_text, tables)
            
        except Exception as e:
            self.logger.error(f"❌ Parse error for page {page_num}: {e}")
            return PageResponse(page_num, response_text, [])

    def _create_table_from_content(self, title_line: str, content: str, table_id: int, page_num: int) -> Optional[ExtractedTable]:
        """Create ExtractedTable from parsed content"""
        try:
            # Extract title
            title = "unknown_table"
            if ':' in title_line:
                title_part = title_line.split(':', 1)[1].strip().rstrip('}')
                title = title_part.lower().replace(' ', '_').replace('-', '_')
            
            # Extract headers and count rows
            lines = content.split('\n')
            headers = []
            row_count = 0
            
            for line in lines:
                if '|' in line and '---' not in line:
                    if not headers:
                        headers = [h.strip() for h in line.split('|')[1:-1]]
                    else:
                        row_count += 1
            
            if not headers:
                return None
            
            return ExtractedTable(
                table_id=table_id,
                title=title,
                markdown_content=content,
                column_headers=headers,
                row_count=row_count,
                column_count=len(headers)
            )
            
        except Exception as e:
            self.logger.error(f"❌ Error creating table: {e}")
            return None

    async def _phase2_bulletproof_sequential_merging(self, pdf_record: PDF, page_responses: List[PageResponse]) -> int:
        """PHASE 2: BULLETPROOF sequential merging with proper tracking"""
        
        total_inserted = 0
        
        self.logger.info(f"🔗 PHASE 2: Starting BULLETPROOF sequential processing...")
        
        for i, current_page in enumerate(page_responses):
            try:
                page_num = current_page.page_number
                tables = current_page.tables
                
                if not tables:
                    continue
                
                self.logger.info(f"🔗 PHASE 2: Processing page {page_num} with {len(tables)} tables...")
                
                # STEP 1: Add all tables except the last one to database
                for j in range(len(tables) - 1):
                    table = tables[j]
                    await self._insert_table_to_database(pdf_record.id, table, page_num, page_num)
                    total_inserted += 1
                    self.logger.info(f"✅ PHASE 2: Added '{table.title}' to database from page {page_num}")
                
                # STEP 2: Handle the last table
                last_table = tables[-1]
                
                # Get next page (immediate next only)
                next_page = None
                if i + 1 < len(page_responses):
                    next_candidate = page_responses[i + 1]
                    if next_candidate.page_number == page_num + 1:
                        next_page = next_candidate
                
                if next_page and next_page.tables:
                    # STEP 3: Check merge with first table of next page
                    first_table_next = next_page.tables[0]
                    
                    self.logger.info(f"🔍 PHASE 2: Checking merge between page {page_num} last table and page {next_page.page_number} first table")
                    
                    # BULLETPROOF merge decision
                    merge_result = await self._bulletproof_merge_decision(last_table, first_table_next, page_num)
                    
                    if merge_result["merged"]:
                        # Merged - replace first table of next page
                        merged_table = merge_result["table"]
                        next_page.tables[0] = merged_table
                        self.logger.info(f"✅ PHASE 2: MERGED - will add merged table when processing page {next_page.page_number}")
                    else:
                        # Not merged - add last table
                        await self._insert_table_to_database(pdf_record.id, last_table, page_num, page_num)
                        total_inserted += 1
                        self.logger.info(f"✅ PHASE 2: SEPARATE - added '{last_table.title}' to database")
                else:
                    # No next page - add last table
                    await self._insert_table_to_database(pdf_record.id, last_table, page_num, page_num)
                    total_inserted += 1
                    self.logger.info(f"✅ PHASE 2: Added final '{last_table.title}' to database")
                
                # Update progress
                pdf_record.tables_processed = total_inserted
                await pdf_record.save()
                        
            except Exception as e:
                self.logger.error(f"❌ PHASE 2: Error processing page {current_page.page_number}: {e}")
                continue
        
        return total_inserted

    async def _bulletproof_merge_decision(self, table1: ExtractedTable, table2: ExtractedTable, current_page: int) -> Dict[str, Any]:
        """LLM ONLY DECISION - No automatic merging"""
        try:
            # ONLY LLM DECIDES - NO AUTOMATIC LOGIC
            prompt = f"""
INTELLIGENT TABLE MERGE DECISION

TABLE 1 (page {current_page}):
Title: {table1.title}
Headers: {table1.column_headers}
Row count: {table1.row_count}
First 3 data rows:
{self._get_first_rows(table1.markdown_content, 3)}

TABLE 2 (page {current_page + 1}):
Title: {table2.title}
Headers: {table2.column_headers}
Row count: {table2.row_count}
First 3 data rows:
{self._get_first_rows(table2.markdown_content, 3)}

TASK: Analyze these two tables carefully. Are they the same logical table continuing across pages?

If YES (should be merged into one continuous table):
Return exactly: MERGE

If NO (they are different tables that should stay separate):
Return exactly: SEPARATE

Look at the data patterns, content, and structure. Make your decision based on whether table 2 is a logical continuation of table 1.

Decision:
"""
            
            client = self._get_next_client()
            
            response = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: client.models.generate_content(
                        model="gemini-2.5-flash-preview-04-17",
                        contents=[prompt]
                    )
                ),
                timeout=25.0
            )
            
            if response and hasattr(response, 'text'):
                decision = str(response.text).strip().upper()
                
                if "MERGE" in decision:
                    # LLM DECIDED MERGE - Do PERFECT merging
                    merged_table = self._perfect_merge_tables(table1, table2, current_page)
                    self.logger.info(f"🔗 PHASE 2: LLM decided MERGE - {merged_table.row_count} total rows")
                    return {"merged": True, "table": merged_table}
                else:
                    # LLM DECIDED SEPARATE - Respect it completely
                    self.logger.info(f"↔️ PHASE 2: LLM decided SEPARATE - keeping tables independent")
                    return {"merged": False}
            
            # LLM failed to respond - default to SEPARATE
            self.logger.warning(f"⚠️ PHASE 2: LLM failed to respond - defaulting to SEPARATE")
            return {"merged": False}
            
        except Exception as e:
            self.logger.error(f"❌ PHASE 2: LLM merge error: {e}")
            # On error, default to SEPARATE
            return {"merged": False}

    def _perfect_merge_tables(self, table1: ExtractedTable, table2: ExtractedTable, current_page: int) -> ExtractedTable:
        """PERFECT table merging that produces correct results"""
        try:
            lines1 = table1.markdown_content.strip().split('\n')
            lines2 = table2.markdown_content.strip().split('\n')
            
            # PERFECT DATA EXTRACTION from table2
            table2_data_rows = []
            found_header = False
            found_separator = False
            
            for line in lines2:
                line = line.strip()
                if not line or '|' not in line:
                    continue
                    
                if not found_header:
                    # Skip the header row
                    found_header = True
                    continue
                elif not found_separator and ('---' in line or '--' in line):
                    # Skip the separator row
                    found_separator = True
                    continue
                else:
                    # This is actual data - add it
                    table2_data_rows.append(line)
            
            # PERFECT MERGE: table1 complete + clean data from table2
            merged_lines = lines1 + table2_data_rows
            merged_content = '\n'.join(merged_lines)
            
            # PERFECT PAGE TRACKING
            actual_start_page = table1.merge_start_page if table1.merge_start_page is not None else current_page
            
            # PERFECT MERGED TABLE
            merged_table = ExtractedTable(
                table_id=table1.table_id,
                title=table1.title,  # Keep original title
                markdown_content=merged_content,
                column_headers=table1.column_headers,  # Keep original headers
                row_count=table1.row_count + table2.row_count,
                column_count=table1.column_count,
                merge_start_page=actual_start_page
            )
            
            self.logger.info(f"🔗 PERFECT MERGE: {table1.row_count} + {table2.row_count} = {merged_table.row_count} rows (pages {actual_start_page}-{current_page + 1})")
            return merged_table
            
        except Exception as e:
            self.logger.error(f"❌ Perfect merge error: {e}")
            # Fallback: return table1 unchanged
            return table1


    def _parse_llm_merge_response_intelligent(self, response_text: str, table1: ExtractedTable, table2: ExtractedTable, current_page: int) -> Dict[str, Any]:
        """INTELLIGENT parsing of LLM merge response"""
        try:
            lines = response_text.strip().split('\n')
            
            if lines and lines[0].strip().upper() == "MERGE":
                self.logger.info(f"🔗 PHASE 2: LLM decided MERGE - parsing response")
                
                # Extract merged table from LLM response
                title = "merged_table"
                table_lines = []
                found_title = False
                
                for line in lines[1:]:
                    line = line.strip()
                    if line.startswith('{') and line.endswith('}') and not found_title:
                        title = line.strip('{}').lower().replace(' ', '_')
                        found_title = True
                    elif '|' in line:
                        table_lines.append(line)
                
                if len(table_lines) >= 2:  # At least header + separator
                    # Use LLM-generated merged table
                    merged_content = '\n'.join(table_lines)
                    
                    # Extract headers from LLM response
                    headers = []
                    if table_lines:
                        header_line = table_lines[0]
                        headers = [h.strip() for h in header_line.split('|')[1:-1]]
                    
                    # Count data rows (exclude header and separator)
                    data_rows = len([line for line in table_lines[2:] if '|' in line and '---' not in line])
                    
                    merged_table = ExtractedTable(
                        table_id=table1.table_id,
                        title=title,
                        markdown_content=merged_content,
                        column_headers=headers,
                        row_count=data_rows,
                        column_count=len(headers),
                        merge_start_page=table1.merge_start_page if table1.merge_start_page is not None else current_page
                    )
                    
                    self.logger.info(f"🔗 PHASE 2: LLM MERGE successful - {merged_table.row_count} rows")
                    return {"merged": True, "table": merged_table}
                else:
                    # LLM response was garbage - fallback to simple merge
                    self.logger.warning(f"⚠️ PHASE 2: LLM merge response was garbage - using simple merge")
                    merged_table = self._bulletproof_merge_tables(table1, table2, current_page)
                    return {"merged": True, "table": merged_table}
            
            else:
                # LLM said SEPARATE or something else
                self.logger.info(f"↔️ PHASE 2: LLM decided SEPARATE")
                return {"merged": False}
                
        except Exception as e:
            self.logger.error(f"❌ PHASE 2: Parse LLM response error: {e}")
            # Fallback - if we can't parse, try simple merge
            merged_table = self._bulletproof_merge_tables(table1, table2, current_page)
            return {"merged": True, "table": merged_table}

    async def _insert_table_to_database(self, pdf_id: PyObjectId, table: ExtractedTable, start_page: int, end_page: int):
        """FIXED: Insert table with proper page tracking"""
        try:
            # FIXED: Proper None checking instead of broken getattr
            actual_start = table.merge_start_page if table.merge_start_page is not None else start_page
            
            table_record = Table(
                pdf_id=pdf_id,
                start_page=actual_start,
                end_page=end_page,
                table_number=table.table_id,
                table_title=table.title,
                markdown_content=table.markdown_content,
                column_count=table.column_count,
                row_count=table.row_count
            )
            
            await table_record.insert()
            self.logger.info(f"💾 FIXED INSERT: '{table.title}' (pages {actual_start}-{end_page}, {table.row_count} rows)")
            
        except Exception as e:
            self.logger.error(f"❌ Database error: {e}")
            raise

    def _bulletproof_merge_tables(self, table1: ExtractedTable, table2: ExtractedTable, start_page: int) -> ExtractedTable:
        """BULLETPROOF table merging"""
        try:
            lines1 = table1.markdown_content.strip().split('\n')
            lines2 = table2.markdown_content.strip().split('\n')
            
            # Find data rows in table2 (skip header and separator)
            data_rows = []
            header_passed = False
            separator_passed = False
            
            for line in lines2:
                if '|' in line:
                    if not header_passed:
                        header_passed = True
                        continue
                    elif not separator_passed and '---' in line:
                        separator_passed = True
                        continue
                    else:
                        data_rows.append(line)
            
            # Merge: table1 + data rows from table2
            merged_lines = lines1 + data_rows
            merged_content = '\n'.join(merged_lines)
            
            merged_table = ExtractedTable(
                table_id=table1.table_id,
                title=table1.title,
                markdown_content=merged_content,
                column_headers=table1.column_headers,
                row_count=table1.row_count + table2.row_count,
                column_count=table1.column_count,
                merge_start_page=table1.merge_start_page if table1.merge_start_page is not None else start_page
            )
            
            self.logger.info(f"🔗 PHASE 2: BULLETPROOF merge - {merged_table.row_count} total rows")
            return merged_table
            
        except Exception as e:
            self.logger.error(f"❌ PHASE 2: Merge error: {e}")
            return table1

    def _get_first_rows(self, markdown_content: str, num_rows: int) -> str:
        """Get first N data rows from markdown"""
        lines = markdown_content.split('\n')
        data_lines = []
        
        for line in lines:
            if '|' in line and '---' not in line:
                data_lines.append(line)
        
        # Skip header, take first num_rows data rows
        if len(data_lines) > 1:
            return '\n'.join(data_lines[1:num_rows+1])
        return "No data"

    async def _cleanup_temp_images(self, images_folder: str):
        """Cleanup temporary images"""
        try:
            if os.path.exists(images_folder):
                shutil.rmtree(images_folder)
                self.logger.info("🧹 Cleaned up temporary images")
        except Exception as e:
            self.logger.error(f"Cleanup error: {e}")


# Background task launcher
async def extract_tables_background(pdf_id: str):
    """Launch BULLETPROOF table extraction"""
    extractor = BackgroundTableExtractor()
    return await extractor.extract_tables_for_pdf(pdf_id)
